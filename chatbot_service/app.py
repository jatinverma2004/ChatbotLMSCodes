from fastapi import FastAPI, Query
from io import BytesIO
from pypdf import PdfReader
from docx import Document
import requests
import json
import threading
import sqlite3
import re
import os
import base64
import random
from datetime import datetime

from chat_evaluator import evaluate_answer
from evaluation_db import insert_record, init_db

# ✅ GROQ IMPORT
from groq import Groq


# ================= DATABASE SETUP =================

DB_PATH = "evaluation.db"

def ensure_dashboard_table():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS evaluations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        question TEXT,
        answer TEXT,
        accuracy REAL,
        precision REAL,
        recall REAL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    """)
    conn.commit()
    conn.close()


init_db()
ensure_dashboard_table()

app = FastAPI(title="Employee Skill Chatbot")
from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ================= CONFIG =================

MCP_BASE_URL = "http://127.0.0.1:8100"

OLLAMA_URL   = "http://127.0.0.1:11434/api/generate"
OLLAMA_MODEL = "phi3:latest"

# ✅ GROQ CLIENT
GROQ_API_KEY = "Your api key here"
client = Groq(api_key=GROQ_API_KEY)
print("✅ GROQ CLIENT INITIALIZED")

SOP_CACHE     = {}   # full text cache  {doc_name: full_text}
SNIPPET_CACHE = {}   # snippet cache    {doc_name: first_300_chars}

# ================= FILE UPLOAD DIRECTORIES =================
# Folders where mcp_v3_server saves uploaded user files
UPLOAD_DIRS = [
    os.path.join(os.path.dirname(__file__), "..", "mcp_v3_server", "file_storage", "snapshots"),
    os.path.join(os.path.dirname(__file__), "..", "mcp_v3_server", "file_storage", "user_uploads"),
    os.path.join(os.path.dirname(__file__), "..", "mcp_v3_server", "sop_storage"),
]

IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}

# ================= FILE HELPERS =================

def find_uploaded_file(filename: str) -> str | None:
    """Search upload dirs for the file. Handles both exact name and unique names like uid_timestamp_original.png"""
    for directory in UPLOAD_DIRS:
        if not os.path.isdir(directory):
            continue
        for root, _, files in os.walk(directory):
            for f in files:
                if f == filename or f.endswith("_" + filename):
                    return os.path.join(root, f)
    return None


def file_to_base64(path: str) -> tuple[str, str]:
    """Return (base64_string, media_type) for an image file."""
    ext = os.path.splitext(path)[1].lower()
    media_map = {
        '.png':  'image/png',
        '.jpg':  'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif':  'image/gif',
        '.bmp':  'image/bmp',
        '.webp': 'image/webp',
    }
    media_type = media_map.get(ext, 'image/png')
    with open(path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8'), media_type


def extract_doc_text_from_path(path: str, filename: str) -> str:
    """Extract text from PDF / DOCX / XLSX / CSV / TXT on disk."""
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext == '.pdf':
            try:
                reader = PdfReader(path)
                text = "\n".join(p.extract_text() or "" for p in reader.pages[:10]).strip()
                return text if text else "[PDF appears to be image-based — no selectable text]"
            except Exception as e:
                return f"[Could not read PDF: {e}]"

        elif ext in ('.docx', '.doc'):
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif ext in ('.xlsx', '.xls'):
            try:
                from openpyxl import load_workbook
                wb = load_workbook(path)
                lines = []
                for ws in wb.worksheets:
                    lines.append(f"[Sheet: {ws.title}]")
                    for row in ws.iter_rows(max_row=30, values_only=True):
                        r = " | ".join(str(c) for c in row if c is not None)
                        if r.strip():
                            lines.append(r)
                return "\n".join(lines)
            except Exception as e:
                return f"[Could not read Excel: {e}]"

        elif ext == '.csv':
            import csv
            lines = []
            with open(path, encoding='utf-8', errors='ignore') as f:
                for i, row in enumerate(csv.reader(f)):
                    if i >= 30:
                        break
                    lines.append(" | ".join(row))
            return "\n".join(lines)

        elif ext == '.txt':
            with open(path, encoding='utf-8', errors='ignore') as f:
                return f.read()[:4000]

        return f"[Unsupported file type: {ext}]"

    except Exception as e:
        return f"[Error reading file: {e}]"


def parse_uploaded_file_from_message(message: str) -> tuple[str, str]:
    """
    If message contains [User uploaded: filename], extract:
    - filename
    - cleaned message (tag stripped out)
    Returns (filename, clean_message). filename is "" if no tag found.
    """
    match = re.search(r'\[User uploaded:\s*(.+?)\]', message)
    if not match:
        return "", message

    filename    = match.group(1).strip()
    clean_msg   = re.sub(r'\[User uploaded:.*?\]', '', message)
    clean_msg   = re.sub(r'File preview:.*', '', clean_msg, flags=re.DOTALL).strip()
    if not clean_msg:
        clean_msg = "Please analyse this file and describe what it contains."
    return filename, clean_msg

# ================= SOP SANITIZER =================

def clean_sop_text(text: str) -> str:
    if not text:
        return ""
    blocked_patterns = [
        "your task", "instruction:", "###", "```",
        "generate", "appendix", "solution:", "problem:", "query:"
    ]
    lower = text.lower()
    for p in blocked_patterns:
        if p in lower:
            text = text.replace(p, "")
    return text.strip()

# ================= OUTPUT GUARD =================

def guard_llm_output(text: str) -> str:
    if not text:
        return text
    lower = text.lower()
    dangerous_patterns = [
        "your task:",
        "generate two more constraints",
        "write an in-depth analysis"
    ]
    for p in dangerous_patterns:
        if p in lower:
            return "⚠️ Unsafe content detected. Please rephrase your question."
    return text

# ================= USER CONTEXT =================

def fetch_user_context(uid: str):
    try:
        r = requests.get(f"{MCP_BASE_URL}/api/context/{uid}")
        return r.json()
    except:
        return {}

# ================= HIVE ROUTER =================

def hive_router(message, context):
    role_code = context.get("user_profile", {}).get("job_role_code", "").upper()
    try:
        r = requests.get(f"{MCP_BASE_URL}/api/sops")
        data = r.json()
        columns = data["columns"]
        rows = data["rows"]
        sops = [dict(zip(columns, row)) for row in rows]
        matched = [s for s in sops if str(s.get("job_role_code", "")).upper() == role_code]
        if matched:
            return {"mode": "SOP", "docs": matched}
        else:
            return {"mode": "GENERAL", "docs": []}
    except:
        return {"mode": "GENERAL", "docs": []}

# ================= SOP TEXT LOADER =================

def fetch_sop_text(sop_name: str) -> str:
    if sop_name in SOP_CACHE:
        return SOP_CACHE[sop_name]
    try:
        url = f"{MCP_BASE_URL}/api/sop/open/{sop_name}"
        r = requests.get(url)
        content_type = r.headers.get("content-type", "")
        if "pdf" in content_type:
            reader = PdfReader(BytesIO(r.content))
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            SOP_CACHE[sop_name] = text
            return text
        if "wordprocessingml" in content_type:
            doc = Document(BytesIO(r.content))
            text = "\n".join([p.text for p in doc.paragraphs if p.text])
            SOP_CACHE[sop_name] = text
            return text
        text = r.text
        SOP_CACHE[sop_name] = text
        return text
    except:
        return ""


def fetch_sop_snippet(sop_name: str, chars: int = 300) -> str:
    if sop_name in SNIPPET_CACHE:
        return SNIPPET_CACHE[sop_name]
    full = fetch_sop_text(sop_name)
    snippet = full[:chars].replace("\n", " ").strip()
    SNIPPET_CACHE[sop_name] = snippet
    return snippet

# ================= SMART DOC SELECTOR =================

def normalize_doc_name(name: str) -> set:
    name = name.lower()
    name = re.sub(r'\.(pdf|docx|doc|txt)$', '', name)
    name = re.sub(r'[_\-\.]+', ' ', name)
    name = re.sub(r'[^a-z0-9\s]', '', name)
    return set(w for w in name.split() if w)

_STOP_WORDS = {
    'what', 'does', 'the', 'say', 'is', 'in', 'a', 'an', 'are', 'how',
    'do', 'i', 'to', 'for', 'of', 'and', 'or', 'tell', 'me', 'about',
    'please', 'can', 'you', 'give', 'show', 'summarize', 'summary',
    'details', 'detail', 'it', 'its', 'this', 'that', 'with', 'from',
    'have', 'has', 'sops', 'docs', 'available', 'list', 'all', 'get'
}

def normalize_query(query: str) -> set:
    query = query.lower()
    query = re.sub(r'[^a-z0-9\s]', '', query)
    return set(w for w in query.split() if w and w not in _STOP_WORDS)


def score_doc(doc_name: str, query_words: set) -> int:
    doc_words   = normalize_doc_name(doc_name)
    doc_name_flat = doc_name.lower()
    score = 0
    for qw in query_words:
        if qw in doc_words:
            score += 100
        elif qw in doc_name_flat:
            score += 50
    return score


def select_relevant_docs(docs, user_message, top_n=3):
    query_words = normalize_query(user_message)
    print(f"🔍 Query words: {query_words}")

    scored = []
    for d in docs:
        s = score_doc(d.get("doc_name", ""), query_words)
        print(f"  {d.get('doc_name')} → score {s}")
        scored.append((s, d))

    scored.sort(key=lambda x: -x[0])
    top_score = scored[0][0] if scored else 0

    if top_score > 0:
        selected = [d for _, d in scored[:top_n]]
    else:
        selected = [d for _, d in scored]

    print(f"✅ Selected: {[d.get('doc_name') for d in selected]}")
    return selected

# ================= DOC INVENTORY BUILDER =================

def build_doc_inventory(docs) -> str:
    lines = ["ASSIGNED SOP DOCUMENTS FOR THIS EMPLOYEE:\n"]
    for i, d in enumerate(docs, 1):
        name = d.get("doc_name", "Unknown")
        snippet = fetch_sop_snippet(name, chars=250)
        display_name = re.sub(r'\.(pdf|docx|doc|txt)$', '', name)
        display_name = display_name.replace("_", " ").replace("-", " ").strip()
        lines.append(f"{i}. [{display_name}] — {snippet}")
    return "\n".join(lines)

# ================= PROMPT =================

def build_prompt(context, sop_text, doc_inventory, user_message, mode):
    user = context.get("user_profile", {})
    role = user.get("job_role_text", "Employee")

    if mode == "SOP":
        return f"""
You are an intelligent Employee Assistant helping employees understand company SOPs.

Instructions:
- Answer clearly and in detail (not just one line)
- Explain concepts in simple language
- If acronym is asked → expand + explain
- Use SOP context strictly (do NOT hallucinate)
- Add examples where useful
- If asked "what documents are available" or "what SOPs exist" → list them from the DOC INVENTORY below
- If a specific doc is asked about → answer from DETAILED SOP CONTEXT
- If answer not found in context → say "Not found in SOP"

ROLE:
{role}

{doc_inventory}

DETAILED SOP CONTEXT (full content of most relevant document(s)):
{sop_text}

USER QUESTION:
{user_message}

FINAL ANSWER:
"""

    return f"""
You are a helpful enterprise assistant.

Answer clearly and explain properly.

ROLE:
{role}

QUESTION:
{user_message}

ANSWER:
"""

# ================= LLM CALL (text only) =================

def call_groq_text(prompt: str) -> str:
    """Standard text-only Groq call — same as before."""
    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("GROQ TEXT ERROR:", e)
        return f"🔥 ERROR: {str(e)}"

# Keep old name as alias so nothing else breaks
call_ollama = call_groq_text

# ================= LLM CALL (vision — image + text) =================

def call_groq_vision(image_b64: str, media_type: str, user_question: str, system_context: str = "") -> str:
    """
    Send an image + question to a vision-capable Groq model.
    Uses meta-llama/llama-4-scout-17b-16e-instruct which supports image_url content blocks.
    """
    try:
        messages = []

        if system_context:
            messages.append({"role": "system", "content": system_context})

        messages.append({
            "role": "user",
            "content": [
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{media_type};base64,{image_b64}"
                    }
                },
                {
                    "type": "text",
                    "text": user_question if user_question else "What does this image show? Describe it in detail."
                }
            ]
        })

        response = client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=messages,
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content.strip()

    except Exception as e:
        print("GROQ VISION ERROR:", e)
        return f"🔥 Vision ERROR: {str(e)}"

# ================= EVALUATION =================

def run_evaluation(uid, message, sop_text, answer):
    ev = evaluate_answer(message, sop_text, answer)
    if ev is None:
        print("❌ Evaluation failed, skipping")
        return
    try:
        insert_record({
            "uid": uid,
            "question": message,
            "retrieved_context": sop_text[:2000],
            "answer": answer,
            "accuracy": ev["accuracy"],
            "precision": ev["precision"],
            "recall": ev["recall"],
            "verdict": ev["verdict"],
            "reason": ev["reason"]
        })
    except Exception as e:
        print("DB INSERT ERROR:", e)

# ================= SUGGESTIONS =================

LAST_SUGGESTIONS = {}

def generate_prompt_suggestions(role, docs, user_query=None, uid=None):
    base = [
        f"What are key responsibilities of a {role}?",
        f"What SOPs should a {role} follow daily?",
        f"What are safety guidelines for {role}?",
        f"How does {role} handle compliance issues?",
        f"What are common mistakes a {role} should avoid?"
    ]

    sop_based = []
    for d in docs:
        name = d.get("doc_name", "").replace(".pdf", "").replace(".docx", "")
        sop_based.extend([
            f"Explain {name}",
            f"Key rules from {name}",
            f"Important points in {name}"
        ])

    query_based = []
    if user_query:
        q = user_query.lower()
        if "leave" in q:
            query_based.extend([
                "What is leave approval workflow?",
                "How many leave types exist?",
                "What documents are required for leave?"
            ])
        if "policy" in q:
            query_based.extend([
                "Summarize policy in simple terms",
                "What are critical policy rules?",
                "What happens if policy is violated?"
            ])
        if "safety" in q:
            query_based.extend([
                "List all safety precautions",
                "How to report safety incidents?",
                "What are safety violations?"
            ])

    all_suggestions = base + sop_based + query_based
    random.shuffle(all_suggestions)

    seen   = set()
    unique = []
    for s in all_suggestions:
        if s not in seen:
            unique.append(s)
            seen.add(s)

    if uid:
        last   = LAST_SUGGESTIONS.get(uid, [])
        unique = [u for u in unique if u not in last]

    final = unique[:5]
    if uid:
        LAST_SUGGESTIONS[uid] = final

    return final

# ================= MAIN CHAT ROUTE =================

@app.post("/chat")
def chat(uid: str = Query(...), message: str = Query(...)):

    context = fetch_user_context(uid)

    # ── Check for uploaded file tag ────────────────────────────────────────────
    filename, clean_message = parse_uploaded_file_from_message(message)

    # ── IMAGE → handle with vision model, skip normal SOP flow ────────────────
    if filename:
        ext = os.path.splitext(filename)[1].lower()

        if ext in IMAGE_EXTS:
            file_path = find_uploaded_file(filename)

            if file_path:
                print(f"🖼️  Vision request: {filename}")
                img_b64, media_type = file_to_base64(file_path)

                # Build a light system context so it knows the user's role
                user  = context.get("user_profile", {})
                role  = user.get("job_role_text", "Employee")
                sys_ctx = f"You are an intelligent assistant for a Jio employee. Their role is: {role}. Analyse the image and answer their question helpfully."

                answer = call_groq_vision(img_b64, media_type, clean_message, sys_ctx)
                answer = guard_llm_output(answer)

                # Suggestions still useful after image answer
                route = hive_router(message, context)
                docs  = route.get("docs", [])
                suggestions = generate_prompt_suggestions(role, docs, clean_message, uid)
                suggestion_text = "\n\n💡 You can also ask:\n" + "".join(f"- {s}\n" for s in suggestions)

                threading.Thread(target=run_evaluation, args=(uid, clean_message, "", answer)).start()
                return {"answer": "🖼️ Image received\n\n" + answer + suggestion_text}

            else:
                print(f"⚠️  Image not found on disk: {filename}")
                answer = f"I received your image ({filename}) but couldn't locate it on the server. Please try uploading again."
                return {"answer": answer}

        # ── DOCUMENT FILE → extract text, inject into prompt ──────────────────
        else:
            file_path = find_uploaded_file(filename)
            if file_path:
                print(f"📄 Document upload: {filename}")
                doc_text = extract_doc_text_from_path(file_path, filename)
                # Prepend extracted text to the message so the normal SOP flow can use it
                message = f"{clean_message}\n\n--- UPLOADED FILE: {filename} ---\n{doc_text[:3000]}\n--- END OF FILE ---"
            else:
                print(f"⚠️  Document not found on disk: {filename}")
                message = clean_message

    # ── Normal SOP / General flow (unchanged from original) ───────────────────
    route = hive_router(message, context)
    mode  = route["mode"]
    docs  = route["docs"]

    doc_inventory = ""
    if mode == "SOP" and docs:
        doc_inventory = build_doc_inventory(docs)

    selected_docs = select_relevant_docs(docs, message, top_n=3)

    collected = []
    for d in selected_docs:
        raw   = fetch_sop_text(d["doc_name"])
        clean = clean_sop_text(raw)
        if clean:
            collected.append(f"=== {d['doc_name']} ===\n{clean}")
    sop_text = "\n\n".join(collected)

    MAX_CHARS = 8000
    if len(sop_text) > MAX_CHARS:
        sop_text = sop_text[:MAX_CHARS]

    prompt = build_prompt(context, sop_text, doc_inventory, message, mode)

    answer = call_groq_text(prompt)
    answer = guard_llm_output(answer)

    suggestions = []
    header      = ""
    final       = answer

    if mode == "SOP" and docs:
        header = f"📄 Using {len(docs)} SOP document(s)\n\n"
        role   = context.get("user_profile", {}).get("job_role_text", "Employee")
        suggestions     = generate_prompt_suggestions(role, docs, message, uid)
        suggestion_text = "\n\n💡 You can also ask:\n" + "".join(f"- {s}\n" for s in suggestions)
        final = header + answer + suggestion_text

    print("FINAL:", final[:300])

    threading.Thread(
        target=run_evaluation,
        args=(uid, message, sop_text, answer)
    ).start()

    if suggestions:
        print("SUGGESTIONS:", suggestions)

    return {"answer": final}

# ================= DASHBOARD DATA =================

@app.get("/dashboard/data")
def dashboard_data():
    try:
        conn   = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM evaluations ORDER BY created_at DESC")
        rows    = cursor.fetchall()
        columns = [c[0] for c in cursor.description]
        conn.close()
        return [dict(zip(columns, row)) for row in rows]
    except Exception as e:
        print("Dashboard error:", e)
        return []
